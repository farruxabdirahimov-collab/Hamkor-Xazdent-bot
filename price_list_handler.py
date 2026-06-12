"""
XazDent Price List Handler
PDF / Rasm / Excel / Word fayllardan mahsulotlarni ajratib oladi
O'xshash mahsulotlarni variantlar bilan guruhlaydi
"""

import io
import os
import json
import logging
import asyncio
import aiohttp
import base64

logger = logging.getLogger(__name__)

GROQ_API_KEY  = os.getenv("GROQ_API_KEY")
GROQ_URL      = "https://api.groq.com/openai/v1/chat/completions"
TEXT_MODEL    = "llama-3.3-70b-versatile"
VISION_MODEL  = "meta-llama/llama-4-scout-17b-16e-instruct"

ANTHROPIC_KEY = os.getenv("ANTHROPIC_API_KEY")
ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"


# ============================================================
# FAYL TURINI ANIQLASH
# ============================================================
def detect_file_type(filename: str, mime_type: str = "") -> str:
    name = filename.lower()
    if name.endswith('.pdf'):                        return 'pdf'
    elif name.endswith(('.xlsx', '.xls')):           return 'excel'
    elif name.endswith(('.docx', '.doc')):           return 'word'
    elif name.endswith(('.jpg', '.jpeg', '.png')):   return 'image'
    elif 'pdf' in mime_type:                         return 'pdf'
    elif 'excel' in mime_type or 'sheet' in mime_type: return 'excel'
    elif 'word' in mime_type:                        return 'word'
    elif 'image' in mime_type:                       return 'image'
    return 'unknown'


# ============================================================
# MATN OLISH
# ============================================================
async def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        return "\n".join(p.extract_text() or "" for p in reader.pages).strip()
    except Exception as e:
        logger.error(f"PDF: {e}")
        return ""

async def extract_text_from_excel(file_bytes: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes))
        ws = wb.active
        rows = []
        for row in ws.iter_rows(values_only=True):
            r = [str(c) if c is not None else "" for c in row]
            if any(r): rows.append(" | ".join(r))
        return "\n".join(rows)
    except Exception as e:
        logger.error(f"Excel: {e}")
        return ""

async def extract_text_from_word(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                r = " | ".join(c.text.strip() for c in row.cells)
                if r.strip(): text += "\n" + r
        return text
    except Exception as e:
        logger.error(f"Word: {e}")
        return ""


# ============================================================
# AI PROMPT — VARIANTLAR BILAN GURUHLASH
# ============================================================
PARSE_PROMPT = """Sen dental (stomatologiya) marketplace uchun price-list tahlilchisan.

MUHIM QOIDA — VARIANTLAR:
O'xshash mahsulotlar (bir xil nom, faqat razmer/rang/artikul farqi) — BITTA mahsulot sifatida variantlar bilan guruhla.

Misol:
  Shprits 1ml — 540 so'm
  Shprits 2ml — 540 so'm     ──► BITTA mahsulot "Shprits bir martalik"
  Shprits 5ml — 560 so'm         3 ta variant bilan
  
  GC Fuji A2 — 85000 so'm
  GC Fuji A3 — 85000 so'm    ──► BITTA mahsulot "GC Fuji IX"
  GC Fuji B2 — 95000 so'm        3 ta variant bilan

Faqat STOMATOLOGIYA mahsulotlarini ol. Telefon, manzil, kompaniya ma'lumotlarini OLMA.

Javobni FAQAT JSON array formatida ber — boshqa hech narsa yozma:
[
  {
    "name": "Mahsulot nomi (qisqa, o'zbekcha/ruscha)",
    "price": 85000,
    "unit": "dona",
    "description": "Qisqa tavsif",
    "stock": 999,
    "variants": [
      {
        "size_name": "1ml",
        "article": "SHP-001",
        "stock": 999,
        "price": 540
      }
    ]
  }
]

QOIDALAR:
- Agar mahsulotning faqat 1 ta varianti bo'lsa — variants: [] (bo'sh)
- Agar narx yo'q bo'lsa — price: 0
- article bo'sh bo'lsa — "" (bo'sh string, avtomatik generatsiya qilinadi)
- stock noma'lum bo'lsa — 999
- unit: dona/quti/ml/g/litr (matndan aniqlash)
"""

async def _call_groq_text(prompt: str, max_tokens: int = 4000) -> str | None:
    if not GROQ_API_KEY:
        return None
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                GROQ_URL,
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": TEXT_MODEL,
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": max_tokens,
                    "temperature": 0.1,
                },
                timeout=aiohttp.ClientTimeout(total=60),
            ) as resp:
                data = await resp.json()
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logger.error(f"Groq text: {e}")
        return None

async def _call_groq_vision(image_bytes: bytes, prompt: str) -> str | None:
    if not GROQ_API_KEY:
        return None
    b64 = base64.b64encode(image_bytes).decode()
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                GROQ_URL,
                headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
                json={
                    "model": VISION_MODEL,
                    "messages": [{"role": "user", "content": [
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                        {"type": "text", "text": prompt}
                    ]}],
                    "max_tokens": 4000,
                    "temperature": 0.1,
                },
                timeout=aiohttp.ClientTimeout(total=60),
            ) as resp:
                data = await resp.json()
        if "choices" not in data:
            logger.error(f"Groq vision: {data.get('error',{}).get('message','')}")
            return None
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        logger.error(f"Groq vision: {e}")
        return None

async def _call_anthropic(file_bytes: bytes, file_type: str) -> str | None:
    if not ANTHROPIC_KEY:
        return None
    b64 = base64.b64encode(file_bytes).decode()
    if file_type == 'image':
        content = [
            {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": b64}},
            {"type": "text", "text": PARSE_PROMPT}
        ]
    elif file_type == 'pdf':
        content = [
            {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": b64}},
            {"type": "text", "text": PARSE_PROMPT}
        ]
    else:
        return None
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                ANTHROPIC_URL,
                headers={"x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
                json={"model": "claude-sonnet-4-5", "max_tokens": 4000, "messages": [{"role": "user", "content": content}]},
                timeout=aiohttp.ClientTimeout(total=60),
            ) as resp:
                data = await resp.json()
        return data["content"][0]["text"].strip()
    except Exception as e:
        logger.error(f"Anthropic: {e}")
        return None


def _parse_json(text: str) -> list:
    """JSON ni xavfsiz parse qilish"""
    try:
        text = text.replace("```json", "").replace("```", "").strip()
        # JSON array boshlanishini topish
        start = text.find('[')
        end   = text.rfind(']') + 1
        if start != -1 and end > start:
            text = text[start:end]
        return json.loads(text)
    except Exception as e:
        logger.error(f"JSON parse: {e}")
        return []


def _clean_products(raw_list: list) -> list:
    """Mahsulotlarni tozalash va to'g'rilash"""
    clean = []
    for i, p in enumerate(raw_list):
        if not p.get("name"):
            continue

        # Variantlarni tozalash
        variants = []
        for v in p.get("variants", []):
            size_name = str(v.get("size_name", "")).strip()
            if not size_name:
                continue
            variants.append({
                "size_name": size_name,
                "article":   str(v.get("article", "")),
                "stock":     int(v.get("stock", 999)),
                "price":     float(v.get("price", p.get("price", 0))),
            })

        price = float(p.get("price", 0))

        clean.append({
            "product_id":  f"pl_{abs(hash(p['name']))%100000}_{i}",
            "title":       str(p.get("name", ""))[:200],
            "description": str(p.get("description", ""))[:500],
            "price_uzs":   int(price),
            "price_usd":   round(price / 12800, 2),
            "unit":        str(p.get("unit", "dona")),
            "stock":       int(p.get("stock", 999)),
            "variants":    variants,
            "min_order":   1,
            "images":      [],
            "photo_file_ids": [],
            "_source":     "price_list",
            "_has_variants": len(variants) > 0,
        })

    return clean


# ============================================================
# ASOSIY FUNKSIYA
# ============================================================
async def process_price_list(file_bytes: bytes, filename: str, mime_type: str = "") -> dict:
    """
    Price-list faylini qayta ishlaydi.
    Qaytaradi: {"ok": True, "products": [...], "count": N, "source": "..."}
    """
    file_type = detect_file_type(filename, mime_type)
    logger.info(f"Price-list: {filename} ({file_type}), {len(file_bytes)//1024} KB")

    raw_list = []
    source   = ""

    # 1. Anthropic — eng kuchli (rasm + PDF uchun)
    if ANTHROPIC_KEY and file_type in ['pdf', 'image']:
        text = await _call_anthropic(file_bytes, file_type)
        if text:
            raw_list = _parse_json(text)
            source   = "Claude AI"

    # 2. Matn ajratib Groq ga berish
    if not raw_list:
        text_content = ""
        if file_type == 'pdf':
            text_content = await extract_text_from_pdf(file_bytes)
        elif file_type == 'excel':
            text_content = await extract_text_from_excel(file_bytes)
        elif file_type == 'word':
            text_content = await extract_text_from_word(file_bytes)

        if text_content and len(text_content) > 50:
            full_prompt = PARSE_PROMPT + f"\n\nPRICE-LIST:\n{text_content[:6000]}"
            text = await _call_groq_text(full_prompt)
            if text:
                raw_list = _parse_json(text)
                source   = "Groq (matn)"

    # 3. Rasm bo'lsa Vision
    if not raw_list and file_type == 'image':
        text = await _call_groq_vision(file_bytes, PARSE_PROMPT)
        if text:
            raw_list = _parse_json(text)
            source   = "Groq Vision"

    if not raw_list:
        return {"ok": False, "error": "Mahsulotlar topilmadi", "products": []}

    products = _clean_products(raw_list)

    # Statistika
    with_variants    = sum(1 for p in products if p["_has_variants"])
    without_variants = len(products) - with_variants

    logger.info(f"Tayyor: {len(products)} mahsulot, {with_variants} ta variantli ({source})")

    return {
        "ok":              True,
        "products":        products,
        "count":           len(products),
        "with_variants":   with_variants,
        "without_variants": without_variants,
        "source":          source,
    }
